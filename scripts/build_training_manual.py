from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "CareOrbit_Client_Training_and_UAT_Manual.docx"
TEAL = "0F766E"
NAVY = "17324D"
LIGHT_TEAL = "E8F5F3"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None, font_size=8.2):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[index], header, True, WHITE, font_size)
        shade(tbl.rows[0].cells[index], TEAL)
    for row_index, row in enumerate(rows):
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, False, NAVY, font_size)
            if row_index % 2:
                shade(cells[index], LIGHT_GRAY)
    if widths:
        for row in tbl.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return tbl


def heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.05
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(":", 1)
        paragraph.add_run(first + ":").bold = True
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    return paragraph


def steps(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def callout(doc, title, text, fill=LIGHT_TEAL):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(title + "  ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    paragraph.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc):
    doc.add_page_break()


def add_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("CareOrbit  |  Client Training & UAT")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Confidential client training material  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("667085")


def add_uat_table(doc, rows):
    return table(
        doc,
        ["ID", "Role", "Test and expected result", "Result", "Notes"],
        rows,
        widths=[0.45, 1.0, 4.35, 0.75, 1.0],
        font_size=7.6,
    )


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    add_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    for name, size, color in [
        ("Title", 30, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 10.5, NAVY),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(4)

    # Cover
    cover = doc.add_table(rows=2, cols=1)
    cover.style = "Table Grid"
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(cover.cell(0, 0), TEAL)
    cover.cell(0, 0).height = Inches(1.2)
    set_cell_text(cover.cell(0, 0), "CAREORBIT", True, WHITE, 18)
    cover.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade(cover.cell(1, 0), LIGHT_TEAL)
    cover.cell(1, 0).text = ""
    p = cover.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.add_run("Client Training &\nUser Acceptance Testing Manual").bold = True
    p.runs[0].font.name = "Aptos Display"
    p.runs[0].font.size = Pt(27)
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("\n\nOperational Guide • Role Workflows • UAT Sign-off")
    p.runs[-1].font.size = Pt(12)
    p.runs[-1].font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run("\n\nVersion 1.0  |  4 July 2026")
    p.runs[-1].font.size = Pt(10)
    p.runs[-1].font.color.rgb = RGBColor.from_string("52606D")
    doc.add_paragraph()
    callout(
        doc,
        "Live system",
        "https://janani-careorbit.vercel.app/",
        LIGHT_BLUE,
    )
    body(
        doc,
        "Use this manual to train staff, perform controlled acceptance testing, and record client approval before routine clinical use.",
    )
    body(
        doc,
        "Patient information is confidential. Use test patients during training and store exported PDF, Excel and JSON files only in approved locations.",
    )

    page_break(doc)
    heading(doc, "1. How To Use This Manual", 1)
    body(
        doc,
        "Training lead: demonstrate each workflow once, then ask the staff member to repeat it using a test patient. Complete the matching UAT case and record Pass or Fail.",
    )
    heading(doc, "Training sequence", 2)
    steps(
        doc,
        [
            "Create or confirm staff accounts and assign the correct roles.",
            "Use one clearly named test patient for the end-to-end demonstration.",
            "Complete registration, appointment, queue, consultation, prescription, lab, pharmacy and billing in that order.",
            "Confirm in-app notifications appear for the receiving role.",
            "Download the PDF, Excel and backup files and open each file.",
            "Record any failed UAT case, retest after correction, and obtain client sign-off.",
        ],
    )
    heading(doc, "System scope", 2)
    table(
        doc,
        ["Available now", "Not part of this release"],
        [
            ("Login and role access", "Inpatient admission and bed management"),
            ("Patient registration and history", "Surgery and operation theatre workflows"),
            ("Appointments, doctor assignment and queue", "Radiology image or DICOM management"),
            ("EMR timeline and prescriptions", "Telemedicine video consultations"),
            ("Lab, pharmacy and billing", "Facial recognition"),
            ("In-app notifications, audit and exports", "Automated WhatsApp/SMS delivery"),
        ],
        widths=[3.35, 3.35],
    )
    callout(
        doc,
        "Important",
        "In-app notifications are active. External WhatsApp/SMS delivery is not included unless a messaging provider is configured separately.",
        LIGHT_BLUE,
    )
    callout(
        doc,
        "Acceptance status",
        "Development verification confirms technical behavior. Client UAT remains pending until the client completes the test cases and signs the acceptance section.",
        LIGHT_TEAL,
    )

    heading(doc, "2. Roles And Access", 1)
    table(
        doc,
        ["Role", "Main responsibilities", "Key areas"],
        [
            ("Admin / Hospital Admin", "Setup, oversight, users, audit and backup", "All operational modules"),
            ("Doctor", "Consultation, clinical review, prescriptions and lab orders", "Patients, EMR, Prescriptions, Laboratory"),
            ("Staff / Nurse", "Registration, appointments, queue and patient coordination", "Patients, Appointments, Queue"),
            ("Lab Technician", "Process lab orders and enter results", "Laboratory"),
            ("Pharmacist", "Maintain stock and dispense medicines", "Pharmacy"),
            ("Billing Operator", "Create invoices, record payments and issue receipts", "Billing"),
        ],
        widths=[1.35, 3.15, 2.25],
    )
    callout(
        doc,
        "Access rule",
        "A restricted message is correct when a user opens a module outside the assigned role. Administrators can access operational modules for supervision.",
    )

    page_break(doc)
    heading(doc, "3. Login, Navigation And Notifications", 1)
    heading(doc, "Sign in", 2)
    steps(
        doc,
        [
            "Open https://janani-careorbit.vercel.app/ in Chrome or Edge.",
            "Enter the email address and password supplied by the administrator.",
            "Select Sign in and wait for the dashboard.",
            "Confirm the user name and expected role are shown.",
            "Use the left navigation to open an allowed module.",
        ],
    )
    heading(doc, "Notifications", 2)
    steps(
        doc,
        [
            "Select the bell icon in the application header.",
            "Review unread items and open the related module when action is required.",
            "Mark reviewed items as read.",
        ],
    )
    bullets(
        doc,
        [
            "A newly assigned appointment notifies the assigned doctor.",
            "A prescription notifies pharmacists.",
            "A new lab order notifies lab technicians.",
            "A completed lab result notifies the ordering doctor.",
        ],
    )
    heading(doc, "Sign out", 2)
    body(
        doc,
        "Always use the profile menu and Sign out when leaving a shared computer. Closing the tab alone is not a secure sign-out.",
    )
    heading(doc, "Forgotten password", 2)
    steps(
        doc,
        [
            "On the sign-in page, select Forgot password?.",
            "Enter the account email and select Send reset link.",
            "Open the CareOrbit recovery email and select the secure link.",
            "Enter and confirm a password that satisfies every displayed requirement.",
            "Select Update password, then sign in again using the new password.",
        ],
    )
    callout(
        doc,
        "Expired link",
        "If the recovery page says the link is invalid, expired or already used, request a new link. Never forward recovery emails.",
        LIGHT_BLUE,
    )

    heading(doc, "4. Patient Registration And History", 1)
    heading(doc, "Register a patient", 2)
    steps(
        doc,
        [
            "Open Patients and select Add patient.",
            "Enter the patient identity, date of birth, gender and contact details.",
            "Add an emergency contact when available.",
            "Review spelling and phone number before saving.",
            "Select Save and confirm that the patient appears with a generated MRN.",
        ],
    )
    heading(doc, "View patient history", 2)
    steps(
        doc,
        [
            "Open Patient History.",
            "Search using the patient name or MRN.",
            "Open the matching record.",
            "Review appointments, clinical events, prescriptions, lab activity and billing history.",
        ],
    )
    callout(
        doc,
        "Duplicate prevention",
        "Search by phone number, name and MRN before registering a returning patient.",
        LIGHT_BLUE,
    )

    page_break(doc)
    heading(doc, "5. Appointment And Queue Workflow", 1)
    heading(doc, "Book and assign", 2)
    steps(
        doc,
        [
            "Open Appointments and select New appointment.",
            "Select the patient and doctor.",
            "Choose the date and time, then add the reason for visit.",
            "Save the appointment.",
            "Confirm the appointment is listed and the assigned doctor receives an in-app notification.",
        ],
    )
    heading(doc, "Reschedule or cancel", 2)
    steps(
        doc,
        [
            "Open the appointment action menu.",
            "Choose Reschedule, select the new date/time and save; or choose Cancel and confirm.",
            "Verify the updated time or cancelled status is visible.",
        ],
    )
    heading(doc, "Manage the queue", 2)
    steps(
        doc,
        [
            "Open Queue and locate the patient.",
            "Move the patient through the available waiting and consultation states.",
            "Confirm the queue order and status update without refreshing the browser.",
        ],
    )

    heading(doc, "6. EMR And Prescription Workflow", 1)
    heading(doc, "Review the EMR timeline", 2)
    steps(
        doc,
        [
            "Open EMR Timeline and select the patient.",
            "Review chronological clinical events and confirm the correct MRN.",
            "Use the timeline as reference before entering new clinical information.",
        ],
    )
    heading(doc, "Issue and download a prescription", 2)
    steps(
        doc,
        [
            "Open Prescriptions and select New prescription.",
            "Select the patient and enter the diagnosis.",
            "Add each medicine with dose, frequency, duration and instructions.",
            "Enter clinical advice and select Issue prescription.",
            "Select PDF on the created prescription and open the downloaded file.",
            "Select Excel to download the prescription register for reporting.",
        ],
    )
    callout(
        doc,
        "Clinical check",
        "The prescribing clinician remains responsible for medicine, dose, interactions and patient-specific advice.",
    )

    page_break(doc)
    heading(doc, "7. Laboratory Workflow", 1)
    heading(doc, "Doctor or administrator", 2)
    steps(
        doc,
        [
            "Open Laboratory and select New lab order.",
            "Select the patient, enter the test name and choose Routine or Urgent.",
            "Create the order and confirm an order number is generated.",
        ],
    )
    heading(doc, "Lab technician", 2)
    steps(
        doc,
        [
            "Open Laboratory and locate the order.",
            "Select Update result.",
            "Change the status as work progresses.",
            "Enter the result and reference range.",
            "Set the status to Completed and save.",
            "Confirm the ordering doctor receives a result-ready notification.",
        ],
    )
    heading(doc, "Export", 2)
    body(
        doc,
        "Select Excel to download the current lab order and result register. Open the workbook and confirm the order number, patient, test, status, result and reference range.",
    )

    heading(doc, "8. Pharmacy Workflow", 1)
    heading(doc, "Add stock", 2)
    steps(
        doc,
        [
            "Open Pharmacy and select Add stock.",
            "Enter medicine name, SKU, batch, expiry date, quantity, reorder level and unit price.",
            "Save and confirm the stock row is visible.",
        ],
    )
    heading(doc, "Dispense medicine", 2)
    steps(
        doc,
        [
            "Select Dispense.",
            "Choose the patient and, when available, link the prescription.",
            "Select the medicine and quantity.",
            "Confirm dispensing.",
            "Verify the stock quantity decreases by exactly the dispensed quantity.",
        ],
    )
    body(
        doc,
        "Select Excel to download two sheets: Stock and Dispensations. Low-stock items are marked in the exported data.",
    )

    page_break(doc)
    heading(doc, "9. Billing, Payments And Receipts", 1)
    heading(doc, "Configure invoice branding (administrator)", 2)
    steps(
        doc,
        [
            "Open Hospital Settings.",
            "Enter the hospital identity, legal, contact, address and tax details.",
            "Upload a PNG, JPG or WebP logo no larger than 2 MB.",
            "Enter invoice terms, payment details, footer and authorized signatory.",
            "Save and verify the logo and hospital details appear in invoice preview.",
        ],
    )
    heading(doc, "Create an invoice", 2)
    steps(
        doc,
        [
            "Open Billing & Payments and select New draft.",
            "Select the patient.",
            "Choose catalogue services or add a custom line with quantity, price and tax rate.",
            "Enter discount and notes when applicable.",
            "Confirm the calculated subtotal, tax and total, then save the draft.",
            "Select Preview and compare patient, hospital and financial details.",
            "Select Finalize invoice only after the preview is correct.",
        ],
    )
    body(
        doc,
        "Drafts may be edited or deleted and do not count as patient debt. Finalized invoice financial lines cannot be edited. An unpaid issued invoice may be cancelled only with a reason; paid invoices require a separate refund process.",
    )
    heading(doc, "Record payment", 2)
    steps(
        doc,
        [
            "Select Record payment on an invoice with an outstanding balance.",
            "Enter an amount not greater than the balance.",
            "Choose cash, card, UPI, bank transfer, insurance or other.",
            "Enter a payment reference when applicable and save.",
            "Confirm paid amount, balance and invoice status are updated.",
            "Select the generated receipt number to download the branded payment receipt PDF.",
        ],
    )
    heading(doc, "Download", 2)
    bullets(
        doc,
        [
            "Select PDF to download the patient invoice/receipt.",
            "Select Excel to download the invoice register with total, paid and balance columns.",
        ],
    )

    heading(doc, "10. Audit, Backup And Downloads", 1)
    heading(doc, "Administrator exports", 2)
    steps(
        doc,
        [
            "Open Audit & Backup.",
            "Review the latest activity rows and confirm user, action and record are visible.",
            "Select Export Excel for a readable multi-sheet operational workbook.",
            "Select Export JSON Backup for the structured application-data snapshot.",
            "Open both files and store them in the approved encrypted backup location.",
        ],
    )
    heading(doc, "Where downloaded files go", 2)
    bullets(
        doc,
        [
            "Chrome and Edge normally save files in the Windows Downloads folder.",
            "Use Ctrl+J to open the browser download list.",
            "Expected file types are .pdf, .xlsx and .json.",
            "If the browser asks for permission, allow downloads for janani-careorbit.vercel.app.",
        ],
    )
    callout(
        doc,
        "Backup limitation",
        "An exported JSON file is an application-data snapshot. Supabase platform backups and recovery settings must also be managed by the system owner.",
        LIGHT_BLUE,
    )

    heading(doc, "11. Controlled Data Imports", 1)
    body(
        doc,
        "Open Data Imports. Only panels allowed for the signed-in role are shown. Use the supplied template whenever possible.",
    )
    heading(doc, "Standard import procedure", 2)
    steps(
        doc,
        [
            "Select Template and enter data without formulas or passwords.",
            "Choose a CSV or XLSX file no larger than 5 MB and 2,000 rows.",
            "Map each CareOrbit field to the matching file column.",
            "Review the five-row preview and correct missing required mappings.",
            "For pharmacy updates, enable existing-record updates only after confirming SKU and batch matches.",
            "Select the confirmation checkbox and then Confirm import.",
            "Review imported and skipped totals; download the error report when rows were rejected.",
        ],
    )
    table(
        doc,
        ["Import", "Authorized roles", "Key validation"],
        [
            ("Patients", "Admin, Staff, Nurse", "MRN, phone, email and name/date-of-birth duplicates"),
            ("Pharmacy", "Admin, Pharmacist", "SKU/batch, expiry, quantity and price"),
            ("Billing services", "Admin, Billing Operator", "Unique code, price, tax and active status"),
            ("Appointments", "Admin, Staff, Nurse", "Existing MRN/doctor, date/time and doctor conflict"),
            ("Staff invitations", "Admin", "Email, supported role and secure activation email; no passwords"),
        ],
        widths=[1.35, 1.65, 3.75],
    )
    callout(
        doc,
        "Import safety",
        "A skipped row is not silently corrected. Review its error, correct the source file and import that row again.",
        LIGHT_BLUE,
    )

    page_break(doc)
    heading(doc, "12. User Acceptance Testing", 1)
    body(
        doc,
        "Use test data only. Enter Pass or Fail in the Result column. A failed case must include notes and be repeated after correction.",
    )
    add_uat_table(
        doc,
        [
            ("UAT-01", "All", "Valid login opens dashboard; invalid login shows an error.", "____", ""),
            ("UAT-02", "All", "Allowed navigation opens; restricted module shows restricted access.", "____", ""),
            ("UAT-03", "Staff", "New patient saves once and receives an MRN.", "____", ""),
            ("UAT-04", "Staff", "Patient search returns the correct name and MRN.", "____", ""),
            ("UAT-05", "Staff", "Appointment saves with selected doctor/date/time.", "____", ""),
            ("UAT-06", "Doctor", "Assigned appointment notification appears in-app.", "____", ""),
            ("UAT-07", "Staff", "Rescheduled appointment shows the new time.", "____", ""),
            ("UAT-08", "Staff", "Cancelled appointment shows cancelled status.", "____", ""),
            ("UAT-09", "Staff", "Queue status changes without duplicate entries.", "____", ""),
            ("UAT-10", "Doctor", "Patient EMR timeline opens for the correct MRN.", "____", ""),
        ],
    )
    add_uat_table(
        doc,
        [
            ("UAT-31", "All", "Forgot-password request shows a generic success message.", "____", ""),
            ("UAT-32", "All", "Recovery email opens reset page and accepts a strong new password.", "____", ""),
            ("UAT-33", "All", "Expired/used recovery link is rejected.", "____", ""),
            ("UAT-34", "Admin", "Hospital profile and invoice text save and reload.", "____", ""),
            ("UAT-35", "Admin", "Valid logo uploads; invalid type/oversize logo is rejected.", "____", ""),
            ("UAT-36", "Billing", "Draft saves without becoming outstanding patient debt.", "____", ""),
            ("UAT-37", "Billing", "Draft edit and delete work; finalized lines cannot be changed.", "____", ""),
            ("UAT-38", "Billing", "Preview and invoice PDF show matching branding and totals.", "____", ""),
            ("UAT-39", "Billing", "Cancellation requires a reason and preserves audit history.", "____", ""),
            ("UAT-40", "Billing", "Payment receipt PDF has receipt/invoice numbers and balance.", "____", ""),
        ],
    )
    add_uat_table(
        doc,
        [
            ("UAT-41", "Staff", "Patient template downloads and valid CSV/XLSX imports.", "____", ""),
            ("UAT-42", "Staff", "Duplicate patient is skipped with a row error.", "____", ""),
            ("UAT-43", "Pharmacist", "Pharmacy import skips matching batch without confirmation.", "____", ""),
            ("UAT-44", "Pharmacist", "Confirmed matching batch update changes intended values.", "____", ""),
            ("UAT-45", "Billing", "Service import updates catalogue and service is selectable.", "____", ""),
            ("UAT-46", "Staff", "Appointment import rejects missing MRN/doctor and conflicts.", "____", ""),
            ("UAT-47", "Doctor", "One notification appears for each successful imported appointment.", "____", ""),
            ("UAT-48", "Admin", "Staff import sends activation email without importing passwords.", "____", ""),
            ("UAT-49", "All", "Formula, invalid-file and oversized-file imports are rejected.", "____", ""),
            ("UAT-50", "All", "Unauthorized roles cannot see or execute import/settings actions.", "____", ""),
        ],
    )
    add_uat_table(
        doc,
        [
            ("UAT-11", "Doctor", "Prescription saves all medicine instructions.", "____", ""),
            ("UAT-12", "Doctor", "Prescription PDF downloads and opens correctly.", "____", ""),
            ("UAT-13", "Doctor", "Prescription Excel downloads and contains the new row.", "____", ""),
            ("UAT-14", "Pharmacist", "New prescription notification appears in-app.", "____", ""),
            ("UAT-15", "Doctor", "Lab order saves with generated order number.", "____", ""),
            ("UAT-16", "Lab Tech", "Lab order notification appears in-app.", "____", ""),
            ("UAT-17", "Lab Tech", "Result and reference range save as Completed.", "____", ""),
            ("UAT-18", "Doctor", "Completed lab result notification appears.", "____", ""),
            ("UAT-19", "Lab Tech", "Lab Excel downloads and contains result data.", "____", ""),
            ("UAT-20", "Pharmacist", "Stock item saves with batch, expiry and price.", "____", ""),
        ],
    )
    add_uat_table(
        doc,
        [
            ("UAT-21", "Pharmacist", "Dispensing reduces stock by the entered quantity.", "____", ""),
            ("UAT-22", "Pharmacist", "Pharmacy Excel contains Stock and Dispensations sheets.", "____", ""),
            ("UAT-23", "Billing", "Invoice total matches items, discount and tax.", "____", ""),
            ("UAT-24", "Billing", "Payment reduces balance and updates status.", "____", ""),
            ("UAT-25", "Billing", "Invoice PDF downloads and opens correctly.", "____", ""),
            ("UAT-26", "Billing", "Billing Excel contains total, paid and balance.", "____", ""),
            ("UAT-27", "Admin", "Audit list shows recent tested actions.", "____", ""),
            ("UAT-28", "Admin", "Admin Excel downloads with multiple data sheets.", "____", ""),
            ("UAT-29", "Admin", "JSON backup downloads and opens as valid JSON.", "____", ""),
            ("UAT-30", "All", "Sign out returns user to login and protects dashboard.", "____", ""),
        ],
    )

    page_break(doc)
    heading(doc, "13. Training Checklist", 1)
    table(
        doc,
        ["Training item", "Completed", "Trainer notes"],
        [
            ("Login, navigation and secure sign-out", "☐", ""),
            ("Forgot-password and reset-password recovery", "☐", ""),
            ("Role permissions and restricted access", "☐", ""),
            ("Patient registration and history", "☐", ""),
            ("Appointment, reschedule, cancel and queue", "☐", ""),
            ("EMR review and prescription", "☐", ""),
            ("Laboratory order and result", "☐", ""),
            ("Pharmacy stock and dispensing", "☐", ""),
            ("Billing, payment and receipt", "☐", ""),
            ("Hospital branding and invoice configuration", "☐", ""),
            ("CSV/XLSX mapping, preview, import and error reports", "☐", ""),
            ("Notifications", "☐", ""),
            ("PDF, Excel and JSON downloads", "☐", ""),
            ("Audit and backup handling", "☐", ""),
            ("Privacy and security responsibilities", "☐", ""),
        ],
        widths=[3.7, 1.0, 2.0],
    )
    heading(doc, "Troubleshooting", 2)
    table(
        doc,
        ["Problem", "Action"],
        [
            ("Page appears stuck", "Wait briefly, refresh once, then sign in again. Record the page and time if repeated."),
            ("Download does not appear", "Use Ctrl+J, allow site downloads, and retry the button once."),
            ("Wrong access", "Ask an administrator to verify the assigned role; do not share another user's account."),
            ("Record not visible", "Confirm search spelling/MRN and refresh the module."),
            ("Notification missing", "Confirm the receiving user's role and refresh the notification panel."),
            ("Incorrect clinical or billing data", "Do not create duplicates. Notify the authorized supervisor for correction."),
        ],
        widths=[2.0, 4.7],
    )

    page_break(doc)
    heading(doc, "14. Security And Operating Rules", 1)
    bullets(
        doc,
        [
            "Each staff member must use an individual account.",
            "Never share passwords, administrator tokens or Supabase/Vercel keys.",
            "Use only the role required for the staff member's job.",
            "Verify the patient name and MRN before any clinical, lab, pharmacy or billing action.",
            "Do not store exported patient files on personal devices or unapproved cloud drives.",
            "Sign out from shared computers and report lost credentials immediately.",
            "Run client-approved backups on a documented schedule and test recovery separately.",
            "Never place passwords in staff import files; invitations must use secure activation emails.",
            "Review invoice drafts before finalization because finalized financial lines are locked.",
        ],
    )
    heading(doc, "Client Acceptance", 1)
    body(
        doc,
        "The undersigned confirms that the listed workflows were demonstrated, tested with approved test data, and accepted subject to any recorded exceptions.",
    )
    table(
        doc,
        ["Name / role", "Signature", "Date"],
        [
            ("Client representative:", "", ""),
            ("Training lead:", "", ""),
            ("System administrator:", "", ""),
        ],
        widths=[2.7, 2.5, 1.5],
    )
    heading(doc, "Exceptions or follow-up actions", 2)
    for _ in range(5):
        body(doc, "________________________________________________________________________________")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
